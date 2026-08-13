"""Write the Northside Dental documents that cannot be hand-written.

Word and PDF files are binary, so the corpus keeps the script that produces
them rather than two blobs nobody can review in a diff. Run it from the
repository root, inside the application image so the writers are installed:

    docker compose run --rm app python sample-projects/write_northside_dental_binaries.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas


PROJECT_FOLDER = Path(__file__).resolve().parent / "northside-dental"
REQUIREMENTS_DOCUMENT = "client-requirements-v1.docx"
TESTING_FEEDBACK_DOCUMENT = "testing-feedback-15-jul.pdf"
LEAVE_POLICY_DOCUMENT = "clinic-staff-leave-policy.pdf"

LEFT_MARGIN = 72
FIRST_LINE_TOP = 760
LINE_HEIGHT = 18


def write_client_requirements() -> None:
    document = Document()
    document.add_heading("Northside Dental — Client Requirements v1", level=1)
    document.add_paragraph("Date: 10 June 2026")
    document.add_paragraph("Prepared by: Software Provider Delivery Owner")
    document.add_paragraph("Approved by: Clinic Practice Manager")

    document.add_heading("Scope", level=1)
    document.add_paragraph(
        "The provider will build online appointment booking for the Northside "
        "Dental website, together with the reminder and schedule work listed "
        "below. This document is the written scope the clinic has approved."
    )

    document.add_heading("Requirements", level=1)
    requirements = [
        ["Requirement", "Detail"],
        [
            "Online booking",
            "A patient must be able to book an appointment from the clinic "
            "website without telephoning reception.",
        ],
        [
            "Appointment reminder",
            "The system must send the patient an email reminder before the "
            "appointment.",
        ],
        [
            "Daily schedule screen",
            "Each dentist must see a doctor-wise list of their own "
            "appointments for the day.",
        ],
        [
            "Cancel or reschedule link",
            "Every booking confirmation must carry a link the patient can use "
            "to cancel or reschedule the appointment.",
        ],
    ]
    table = document.add_table(rows=len(requirements), cols=2)
    for row_index, row in enumerate(requirements):
        for cell_index, cell_text in enumerate(row):
            table.cell(row_index, cell_index).text = cell_text

    document.add_heading("Out of scope", level=1)
    document.add_paragraph(
        "Anything not listed under Requirements is out of scope for this "
        "release."
    )
    document.save(str(PROJECT_FOLDER / REQUIREMENTS_DOCUMENT))


def write_testing_feedback() -> None:
    _write_pdf(
        PROJECT_FOLDER / TESTING_FEEDBACK_DOCUMENT,
        [
            [
                "Northside Dental — Testing Feedback",
                "Date: 15 July 2026",
                "Tested by: Clinic Practice Manager",
                "",
                "What we tried",
                "We booked six test appointments from the clinic website and",
                "opened the daily schedule screen for two of our dentists.",
                "",
                "What we found",
                "Online booking works. Every test appointment was saved against",
                "the dentist we chose.",
                "",
                "The email reminder works. Each test patient received the",
                "reminder before the appointment.",
                "",
                "The daily schedule screen shows the wrong day. It opens on",
                "tomorrow's list instead of today's.",
                "",
                "The SMS reminder still does not reach the patient. We are",
                "logging this one as a bug.",
            ]
        ],
    )


def write_leave_policy() -> None:
    _write_pdf(
        PROJECT_FOLDER / LEAVE_POLICY_DOCUMENT,
        [
            [
                "Northside Dental — Staff Leave Policy",
                "Date: 1 April 2026",
                "",
                "Annual leave",
                "Clinical and reception staff receive twenty-four days of paid",
                "annual leave in a calendar year.",
                "",
                "Requesting leave",
                "Leave is requested through the practice manager at least two",
                "weeks before the first day away.",
                "",
                "Sick leave",
                "Staff telephone the practice manager before nine in the",
                "morning on the first day of any sickness absence.",
            ]
        ],
    )


def _write_pdf(path: Path, pages: list[list[str]]) -> None:
    pdf = canvas.Canvas(str(path))
    for lines in pages:
        written = pdf.beginText(LEFT_MARGIN, FIRST_LINE_TOP)
        written.setLeading(LINE_HEIGHT)
        for line in lines:
            written.textLine(line)
        pdf.drawText(written)
        pdf.showPage()
    pdf.save()


if __name__ == "__main__":
    write_client_requirements()
    write_testing_feedback()
    write_leave_policy()
