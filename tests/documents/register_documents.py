from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.examine.examine_register import EXAMINE_PROMPT_MARKER
from app.extract.answer import MEETING_NOTES
from app.match.match_requirements import EXISTING_ROW, MATCH_PROMPT_MARKER, NEW_ROW


EXTRACT_MARKER = "File name: {source_file}"
# Only Match is sent the requirements as JSON, so this marker picks out the
# Match call of the run whose batch holds one named document.
MATCH_BATCH_MARKER = '"source_file": "{source_file}"'
# A first run's Match sees an empty register, so on a first run this row text
# reaches the model only in the Examine call of the project that holds the row.
EXAMINE_ROW_MARKER = '"what_was_asked": "{what_was_asked}"'
# The register Match is shown is empty only before any row is committed, which
# is what tells a first run's Match call apart from a later run's.
EMPTY_REGISTER_MARKER = "Register rows:\n[]"

# Not a secret: it locks a fixture that one test writes and throws away.
FIXTURE_PDF_PASSWORD = "fixture-pdf-lock"
FIRST_LINE_TOP = 720
LEFT_MARGIN = 72
LINE_HEIGHT = 18


def write_meeting_note(folder: Path, source_file: str, requirement: str) -> str:
    """One fabricated meeting note, and the words a citation must be traced to."""
    quote = f"The client asked for {requirement}."
    (folder / source_file).write_text(
        f"# Intake portal — {source_file}\n\n"
        "**Date:** 10 March 2026\n\n"
        "## Discussion\n\n"
        f"{quote}\n",
        encoding="utf-8",
    )
    return quote


def write_pdf(path: Path, pages: list[list[str]]) -> None:
    """One fabricated PDF with a real text layer — one list of lines per page."""
    pdf = canvas.Canvas(str(path))
    for lines in pages:
        written = pdf.beginText(LEFT_MARGIN, FIRST_LINE_TOP)
        written.setLeading(LINE_HEIGHT)
        for line in lines:
            written.textLine(line)
        pdf.drawText(written)
        pdf.showPage()
    pdf.save()


def write_pdf_without_a_text_layer(path: Path) -> None:
    """The shape a scanned page has: marks on the page and no text at all."""
    pdf = canvas.Canvas(str(path))
    pdf.rect(LEFT_MARGIN, 500, 400, 200, fill=1)
    pdf.showPage()
    pdf.save()


def write_encrypted_pdf(path: Path, pages: list[list[str]]) -> None:
    write_pdf(path, pages)
    writer = PdfWriter(clone_from=str(path))
    writer.encrypt(FIXTURE_PDF_PASSWORD)
    with path.open("wb") as encrypted_file:
        writer.write(encrypted_file)


def write_corrupt_pdf(path: Path) -> None:
    """A file the folder accepts as a PDF and no PDF library can open."""
    path.write_bytes(b"%PDF-1.4\nthis file was truncated before it was written")


# The font object reportlab writes for every page's default font, byte for
# byte. Overwriting it with garbage of the same length leaves every later
# object at its original offset, so pypdf's page count (which never resolves
# a page's fonts) still succeeds while pdfminer's parser — reached only once
# text is actually extracted — fails on it.
_PDF_FONT_OBJECT = (
    b"<<\n/BaseFont /Helvetica /Encoding /WinAnsiEncoding /Name /F1 "
    b"/Subtype /Type1 /Type /Font\n>>"
)


def write_pdf_with_unparseable_content(path: Path) -> None:
    """A PDF pypdf can count the pages of, but pdfplumber cannot read the text of."""
    write_pdf(path, [["This page will not be read."]])
    data = path.read_bytes()
    if data.count(_PDF_FONT_OBJECT) != 1:
        raise AssertionError(
            "write_pdf's output no longer matches this fixture's corruption "
            "pattern — reportlab's default PDF layout must have changed"
        )
    garbage = b"<<" + b"$" * (len(_PDF_FONT_OBJECT) - 4) + b">>"
    path.write_bytes(data.replace(_PDF_FONT_OBJECT, garbage))


def write_corrupt_docx(path: Path) -> None:
    """A `.docx` that is not a Word package at all — Word files are zips."""
    path.write_bytes(b"this was renamed to .docx and is not a Word file")


def write_docx(
    path: Path,
    sections: list[tuple[str, list[str]]],
    table_rows: list[list[str]] | None = None,
) -> None:
    """One fabricated Word document of headed paragraphs and an optional table."""
    document = Document()
    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, cell_text in enumerate(row):
                table.cell(row_index, cell_index).text = cell_text
    document.save(str(path))


def write_text_file(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def requirement_extraction_answer(
    summary: str,
    quote: str,
    document_type: str = MEETING_NOTES,
) -> dict[str, Any]:
    """One requirement from a document that states no date of its own."""
    return {
        "document_type": document_type,
        "document_date": None,
        "requirements": [{"summary": summary, "quote": quote}],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def extraction_answer(summary: str, quote: str) -> dict[str, Any]:
    return {
        "document_type": "meeting notes",
        "document_date": {
            "value": "10 March 2026",
            "quote": "**Date:** 10 March 2026",
        },
        "requirements": [{"summary": summary, "quote": quote}],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def unrelated_extraction_answer() -> dict[str, Any]:
    """A document Extract judged unrelated, reporting nothing in any list.

    An unrelated document may fill none of the four lists, so this is the only
    shape one can honestly come back in.
    """
    return extraction_answer_without_requirements() | {"document_type": "unrelated"}


def feedback_extraction_answer(
    observations: list[tuple[str, str, str]],
    date: str = "25 March 2026",
) -> dict[str, Any]:
    """What testing found, as testing feedback: summary, label and quote each."""
    return {
        "document_type": "testing feedback",
        "document_date": {"value": date, "quote": f"**Date:** {date}"},
        "requirements": [],
        "testing_observations": [
            {"summary": summary, "label": label, "quote": quote}
            for summary, label, quote in observations
        ],
        "delivery_evidence": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def handover_answer(
    delivered: list[tuple[str, str]],
    date: str = "20 July 2026",
) -> dict[str, Any]:
    """What a handover summary reports as handed over, and nothing else."""
    return {
        "document_type": "related additional document",
        "document_date": {"value": date, "quote": f"**Date:** {date}"},
        "requirements": [],
        "testing_observations": [],
        "delivery_evidence": [
            {"summary": summary, "quote": quote} for summary, quote in delivered
        ],
        "blockers": [],
        "embedded_instructions": [],
    }


def extraction_answer_without_requirements() -> dict[str, Any]:
    """A document of this engagement that simply asks for nothing."""
    return {
        "document_type": "meeting notes",
        "document_date": None,
        "requirements": [],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def dated_extraction_answer(
    requirements: list[tuple[str, str]],
    document_type: str,
    date: str,
) -> dict[str, Any]:
    """Several requirements out of one document that states its own date."""
    return {
        "document_type": document_type,
        "document_date": {"value": date, "quote": f"**Date:** {date}"},
        "requirements": [
            {"summary": summary, "quote": quote} for summary, quote in requirements
        ],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def match_answer(requirement_count: int) -> dict[str, Any]:
    return {
        "outcomes": [
            {"requirement_index": index, "outcome": "new row", "row_number": None}
            for index in range(requirement_count)
        ]
    }


def match_answer_existing_row(row_number: int) -> dict[str, Any]:
    """Match reporting that the one requirement in this batch is already a row."""
    return {
        "outcomes": [
            {
                "requirement_index": 0,
                "outcome": "existing row",
                "row_number": row_number,
            }
        ]
    }


def match_answer_of(matched_row_numbers: list[int | None]) -> dict[str, Any]:
    """Match's answer per requirement: the row it matched, or None for a new row."""
    return {
        "outcomes": [
            {
                "requirement_index": index,
                "outcome": NEW_ROW if row_number is None else EXISTING_ROW,
                "row_number": row_number,
            }
            for index, row_number in enumerate(matched_row_numbers)
        ]
    }


def extract_marker(source_file: str) -> str:
    return EXTRACT_MARKER.format(source_file=source_file)


def match_marker() -> str:
    return MATCH_PROMPT_MARKER


def match_marker_against_an_empty_register() -> str:
    """The Match call of a run whose project has no committed row yet."""
    return EMPTY_REGISTER_MARKER


def examine_marker() -> str:
    return EXAMINE_PROMPT_MARKER


def examine_marker_for_register_holding(what_was_asked: str) -> str:
    """The first run's Examine call for the project whose register holds this row."""
    return EXAMINE_ROW_MARKER.format(what_was_asked=what_was_asked)


def no_findings_answer() -> dict[str, Any]:
    """Examine reporting an honest empty result for a register with nothing wrong."""
    return {"findings": []}


def match_marker_for_batch_with(source_file: str) -> str:
    return MATCH_BATCH_MARKER.format(source_file=source_file)
