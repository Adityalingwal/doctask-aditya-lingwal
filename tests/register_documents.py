from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.examine.examine_register import EXAMINE_PROMPT_MARKER
from app.extract.answer import MEETING_NOTES
from app.match.match_requirements import MATCH_PROMPT_MARKER


EXTRACT_MARKER = "File name: {source_file}"
# Only Match is sent the requirements as JSON, so this marker picks out the
# Match call of the run whose batch holds one named document.
MATCH_BATCH_MARKER = '"source_file": "{source_file}"'

# Not a secret: it locks a fixture that one test writes and throws away.
FIXTURE_PDF_PASSWORD = "fixture-pdf-lock"
FIRST_LINE_TOP = 720
LEFT_MARGIN = 72
LINE_HEIGHT = 18


def write_meeting_note(folder: Path, source_file: str, requirement: str) -> str:
    """One fabricated meeting note, and the words a citation must be traced to."""
    quote = f"The client asked for {requirement}."
    (folder / source_file).write_text(
        f"# Intake portal — {source_file}\n\n"
        "**Date:** 10 March 2026\n\n"
        "## Discussion\n\n"
        f"{quote}\n",
        encoding="utf-8",
    )
    return quote


def write_pdf(path: Path, pages: list[list[str]]) -> None:
    """One fabricated PDF with a real text layer — one list of lines per page."""
    pdf = canvas.Canvas(str(path))
    for lines in pages:
        written = pdf.beginText(LEFT_MARGIN, FIRST_LINE_TOP)
        written.setLeading(LINE_HEIGHT)
        for line in lines:
            written.textLine(line)
        pdf.drawText(written)
        pdf.showPage()
    pdf.save()


def write_pdf_without_a_text_layer(path: Path) -> None:
    """The shape a scanned page has: marks on the page and no text at all."""
    pdf = canvas.Canvas(str(path))
    pdf.rect(LEFT_MARGIN, 500, 400, 200, fill=1)
    pdf.showPage()
    pdf.save()


def write_encrypted_pdf(path: Path, pages: list[list[str]]) -> None:
    write_pdf(path, pages)
    writer = PdfWriter(clone_from=str(path))
    writer.encrypt(FIXTURE_PDF_PASSWORD)
    with path.open("wb") as encrypted_file:
        writer.write(encrypted_file)


def write_corrupt_pdf(path: Path) -> None:
    """A file the folder accepts as a PDF and no PDF library can open."""
    path.write_bytes(b"%PDF-1.4\nthis file was truncated before it was written")


def write_corrupt_docx(path: Path) -> None:
    """A `.docx` that is not a Word package at all — Word files are zips."""
    path.write_bytes(b"this was renamed to .docx and is not a Word file")


def write_docx(
    path: Path,
    sections: list[tuple[str, list[str]]],
    table_rows: list[list[str]] | None = None,
) -> None:
    """One fabricated Word document of headed paragraphs and an optional table."""
    document = Document()
    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, cell_text in enumerate(row):
                table.cell(row_index, cell_index).text = cell_text
    document.save(str(path))


def write_text_file(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def requirement_extraction_answer(
    summary: str,
    quote: str,
    document_type: str = MEETING_NOTES,
) -> dict[str, Any]:
    """One requirement from a document that states no date of its own."""
    return {
        "document_type": document_type,
        "document_date": None,
        "requirements": [{"summary": summary, "quote": quote}],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def extraction_answer(summary: str, quote: str) -> dict[str, Any]:
    return {
        "document_type": "meeting notes",
        "document_date": {
            "value": "10 March 2026",
            "quote": "**Date:** 10 March 2026",
        },
        "requirements": [{"summary": summary, "quote": quote}],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def unrelated_extraction_answer(summary: str, quote: str) -> dict[str, Any]:
    """A document Extract judged unrelated, which still listed a requirement.

    The Extract node already guards against this shape by forcing the
    requirement count to zero for an unrelated document, so a batch can end
    with nothing found while the stored extraction is not empty.
    """
    return extraction_answer(summary, quote) | {"document_type": "unrelated"}


def extraction_answer_without_requirements() -> dict[str, Any]:
    """A document of this engagement that simply asks for nothing."""
    return {
        "document_type": "meeting notes",
        "document_date": None,
        "requirements": [],
        "testing_observations": [],
        "blockers": [],
        "embedded_instructions": [],
    }


def match_answer(requirement_count: int) -> dict[str, Any]:
    return {
        "outcomes": [
            {"requirement_index": index, "outcome": "new row", "row_number": None}
            for index in range(requirement_count)
        ]
    }


def match_answer_existing_row(row_number: int) -> dict[str, Any]:
    """Match reporting that the one requirement in this batch is already a row."""
    return {
        "outcomes": [
            {
                "requirement_index": 0,
                "outcome": "existing row",
                "row_number": row_number,
            }
        ]
    }


def extract_marker(source_file: str) -> str:
    return EXTRACT_MARKER.format(source_file=source_file)


def match_marker() -> str:
    return MATCH_PROMPT_MARKER


def examine_marker() -> str:
    return EXAMINE_PROMPT_MARKER


def no_findings_answer() -> dict[str, Any]:
    """Examine reporting an honest empty result for a register with nothing wrong."""
    return {"findings": []}


def match_marker_for_batch_with(source_file: str) -> str:
    return MATCH_BATCH_MARKER.format(source_file=source_file)
